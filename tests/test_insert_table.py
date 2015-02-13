__author__ = 'snegovik'

import unittest
from docx.api import Document


class TestAddTable(unittest.TestCase):
    def test_insert_table(self):
        import os

        docx_ = "hello.docx"
        path = os.path.dirname(os.path.realpath(__file__))
        docx_ = os.path.join(path, "test_files", docx_)
        assert (os.path.exists(docx_))
        document = Document(docx_)
        assert (isinstance(document, Document))
        assert (len(document.tables) == 1)

        para = document.paragraphs[0]
        table = document.insert_table(3, 3, para, style=None)
        for row_index in xrange(3):
            for column_index, cell in enumerate(table.row_cells(row_index)):
                cell.text = str((row_index + 1) * (column_index + 1))

        docx_out = docx_.replace("hello.docx", "hello_out.docx")
        document.save(docx_out)

        document_out = Document(docx_out)
        assert (len(document_out.tables) == 1)

        # Another table is inside of first Paragraph
